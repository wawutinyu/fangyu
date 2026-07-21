"""Bundle 内 Coding 手脚 — 供 agent_loop 调用，全部限定在 active workspace。"""
from __future__ import annotations

import json
import re
import subprocess
from pathlib import Path
from typing import Any

from fangyu.engine.workspace import WorkspaceError, get_active_workspace


def _ws():
    ws = get_active_workspace()
    if not ws:
        raise WorkspaceError("无 active workspace — 请先 init_bundle_workspace")
    return ws


def tool_read(path: str = "") -> str:
    return _ws().read(path)


def tool_write(path: str = "", content: str = "") -> str:
    _ws().write(path, content)
    return f"wrote {path} ({len(content)} chars)"


def tool_list(path: str = ".") -> list[str]:
    return _ws().list(path)


def tool_search(pattern: str = "", path: str = ".", max_hits: int = 50) -> list[dict[str, Any]]:
    """在工作区内按正则搜文件内容（grep 别名，兼容旧名）。"""
    return tool_grep(pattern=pattern, path=path, max_hits=max_hits)


def tool_grep(pattern: str = "", path: str = ".", max_hits: int = 50) -> list[dict[str, Any]]:
    """在工作区内按正则搜文件内容。"""
    ws = _ws()
    root = ws.resolve(path)
    if not pattern:
        return []
    rx = re.compile(pattern)
    hits: list[dict[str, Any]] = []
    files = [root] if root.is_file() else sorted(root.rglob("*"))
    for f in files:
        if not f.is_file():
            continue
        if ".fangyu" in f.parts:
            continue
        try:
            text = f.read_text(encoding="utf-8", errors="ignore")
        except OSError:
            continue
        rel = str(f.relative_to(ws.root))
        for i, line in enumerate(text.splitlines(), 1):
            if rx.search(line):
                hits.append({"path": rel, "line": i, "text": line[:200]})
                if len(hits) >= max_hits:
                    return hits
    return hits


def tool_glob(pattern: str = "**/*", path: str = ".") -> list[str]:
    """按 glob 模式列出工作区相对路径（文件）。"""
    import fnmatch

    ws = _ws()
    base = ws.resolve(path)
    pat = (pattern or "**/*").strip() or "**/*"
    out: list[str] = []
    if base.is_file():
        rel = str(base.relative_to(ws.root))
        if fnmatch.fnmatch(rel, pat) or fnmatch.fnmatch(base.name, pat):
            return [rel]
        return []
    # 支持 ** 与简单后缀
    if "**" in pat or "/" in pat:
        files = sorted(base.rglob("*"))
        for f in files:
            if not f.is_file() or ".fangyu" in f.parts:
                continue
            rel = str(f.relative_to(ws.root))
            # 相对 path 根的匹配 + 相对 workspace 根
            try:
                rel_base = str(f.relative_to(base))
            except ValueError:
                rel_base = rel
            if fnmatch.fnmatch(rel, pat) or fnmatch.fnmatch(rel_base, pat) or fnmatch.fnmatch(f.name, pat):
                out.append(rel)
    else:
        for f in sorted(base.rglob(pat)):
            if f.is_file() and ".fangyu" not in f.parts:
                out.append(str(f.relative_to(ws.root)))
    return out[:500]


def _strip_html(html: str) -> str:
    text = re.sub(r"(?is)<script[^>]*>.*?</script>", " ", html)
    text = re.sub(r"(?is)<style[^>]*>.*?</style>", " ", text)
    text = re.sub(r"(?s)<[^>]+>", " ", text)
    text = re.sub(r"&nbsp;", " ", text)
    text = re.sub(r"&amp;", "&", text)
    text = re.sub(r"&lt;", "<", text)
    text = re.sub(r"&gt;", ">", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def tool_webfetch(url: str = "", max_chars: int = 12000) -> dict[str, Any]:
    """拉取 URL 文本（HTML 粗转文本）。工厂原料 webfetch。"""
    import httpx

    u = (url or "").strip()
    if not u:
        raise ValueError("url 为空")
    if not re.match(r"^https?://", u, re.I):
        raise ValueError("仅允许 http/https URL")
    with httpx.Client(timeout=20.0, follow_redirects=True) as client:
        resp = client.get(u, headers={"User-Agent": "fangyu-webfetch/1.0"})
        resp.raise_for_status()
        ctype = (resp.headers.get("content-type") or "").lower()
        raw = resp.text
        final_url = str(resp.url)
        status_code = resp.status_code
    if "html" in ctype or raw.lstrip().startswith("<"):
        body = _strip_html(raw)
    else:
        body = raw
    total_len = len(body)
    truncated = total_len > max_chars
    if truncated:
        body = body[:max_chars] + f"\n…(truncated, total {total_len} chars)"
    return {
        "url": final_url,
        "status_code": status_code,
        "content_type": ctype,
        "text": body,
        "truncated": truncated,
    }


def tool_websearch(query: str = "", max_results: int = 5) -> dict[str, Any]:
    """互联网搜索（DuckDuckGo Instant Answer API，无 Key）。"""
    import httpx
    from urllib.parse import urlencode

    q = (query or "").strip()
    if not q:
        raise ValueError("query 为空")
    n = max(1, min(int(max_results or 5), 10))
    params = urlencode({"q": q, "format": "json", "no_html": "1", "skip_disambig": "1"})
    api = f"https://api.duckduckgo.com/?{params}"
    with httpx.Client(timeout=20.0, follow_redirects=True) as client:
        resp = client.get(api, headers={"User-Agent": "fangyu-websearch/1.0"})
        resp.raise_for_status()
        data = resp.json()
    results: list[dict[str, str]] = []
    abstract = (data.get("AbstractText") or "").strip()
    abstract_url = (data.get("AbstractURL") or "").strip()
    if abstract:
        results.append({
            "title": data.get("Heading") or q,
            "url": abstract_url,
            "snippet": abstract[:500],
        })
    for topic in data.get("RelatedTopics") or []:
        if len(results) >= n:
            break
        if isinstance(topic, dict) and topic.get("Text"):
            results.append({
                "title": (topic.get("Text") or "")[:80],
                "url": str(topic.get("FirstURL") or ""),
                "snippet": (topic.get("Text") or "")[:400],
            })
        elif isinstance(topic, dict) and isinstance(topic.get("Topics"), list):
            for sub in topic["Topics"]:
                if len(results) >= n:
                    break
                if isinstance(sub, dict) and sub.get("Text"):
                    results.append({
                        "title": (sub.get("Text") or "")[:80],
                        "url": str(sub.get("FirstURL") or ""),
                        "snippet": (sub.get("Text") or "")[:400],
                    })
    return {"query": q, "results": results[:n], "source": "duckduckgo"}


def tool_question(prompt: str = "", options: str = "") -> dict[str, Any]:
    """向人类提问（观察结果回灌；调用方应在 done 中转述并等待）。"""
    q = (prompt or "").strip()
    if not q:
        raise ValueError("prompt 为空")
    opts = [o.strip() for o in (options or "").split("|") if o.strip()]
    # 落盘便于观/审计
    try:
        ws = _ws()
        log = ws.root / ".fangyu" / "questions.jsonl"
        log.parent.mkdir(parents=True, exist_ok=True)
        import json as _json
        import time
        with log.open("a", encoding="utf-8") as f:
            f.write(_json.dumps({
                "ts": time.time(),
                "question": q,
                "options": opts,
            }, ensure_ascii=False) + "\n")
    except Exception:
        pass
    return {
        "status": "needs_user",
        "question": q,
        "options": opts,
        "hint": "请把问题转述给用户；取得答复前不要臆造答案。",
    }


def tool_apply_patch(path: str = "", old: str = "", new: str = "") -> str:
    """简单字符串替换补丁（单文件）。"""
    ws = _ws()
    cur = ws.read(path)
    if old not in cur:
        raise ValueError(f"patch old 片段未找到: {path}")
    ws.write(path, cur.replace(old, new, 1))
    return f"patched {path}"


_SHELL_DENY = re.compile(
    r"(rm\s+-rf\s+/|sudo\s+|mkfs|dd\s+if=|:\(\)\s*\{|curl\s+[^\n]*\|\s*sh)",
    re.I,
)
# S0-B2：禁止依赖 shell 元字符（管道/重定向/替换），强制 argv
_SHELL_META = re.compile(r"[|;&$`<>\n\\]")


def tool_shell(
    command: str = "",
    timeout_sec: float = 30,
    confirm: bool = False,
    approval_id: str = "",
) -> dict[str, Any]:
    """在 workspace 根目录执行命令（argv 列表，无 shell=True）。

    策略（contextvar shell_policy）：
      - deny: 一律拒绝
      - ask: 非只读命令需人审批准（approval_id），否则返回 needs_approval
      - allow: 仅受危险命令黑名单约束
    """
    import os
    import shlex

    from fangyu.engine.approval_queue import (
        consume_shell_approval,
        enqueue_shell_approval,
        get_approval,
    )
    from fangyu.engine.shell_policy import get_shell_policy, shell_needs_confirm

    cmd = (command or "").strip()
    if not cmd:
        raise ValueError("empty command")
    policy = get_shell_policy()
    if policy == "deny":
        raise PermissionError("shell 策略为 deny，禁止执行")
    if _SHELL_DENY.search(cmd):
        raise PermissionError(f"命令被策略拒绝: {cmd[:80]}")
    if _SHELL_META.search(cmd):
        raise PermissionError(
            "禁止 shell 元字符（|;&$`<>\\）；请用单命令 argv，或用 write/apply_patch 改文件"
        )
    if shell_needs_confirm(cmd):
        aid = (approval_id or "").strip()
        if confirm and aid and consume_shell_approval(aid, cmd):
            pass  # 已消费批准，继续执行
        else:
            # 复用未决/已批同一命令的 pending/approved 项，避免刷屏
            existing = None
            if aid:
                existing = get_approval(aid)
            if not existing or existing.get("command") != cmd or existing.get("status") not in (
                "pending", "approved",
            ):
                existing = enqueue_shell_approval(cmd)
            return {
                "status": "needs_approval",
                "approval_id": existing["id"],
                "command": cmd,
                "policy": "ask",
                "approval_status": existing.get("status"),
                "hint": (
                    "非只读 shell，已进入人审队列。"
                    "请在 Studio「运维 → 人审」批准；"
                    "批准后可用同一 approval_id 且 confirm=true 重试，"
                    "或在人审面板点「批准并执行」。"
                ),
            }
    try:
        argv = shlex.split(cmd, posix=(os.name != "nt"))
    except ValueError as e:
        raise ValueError(f"无法解析命令: {e}") from e
    if not argv:
        raise ValueError("empty command")
    ws = _ws()
    proc = subprocess.run(
        argv,
        shell=False,
        cwd=str(ws.root),
        capture_output=True,
        text=True,
        timeout=timeout_sec,
    )
    return {
        "exit_code": proc.returncode,
        "stdout": (proc.stdout or "")[:8000],
        "stderr": (proc.stderr or "")[:4000],
        "confirmed": bool(confirm),
        "approval_id": (approval_id or "").strip() or None,
    }


def coding_toolbelt() -> dict[str, Any]:
    """agent_loop 用的默认编码工具表（含工厂 P0 原料）。"""
    from fangyu.core.skill_pack import tool_skill_load
    from fangyu.engine.browser_tool import (
        tool_browser_click,
        tool_browser_open,
        tool_browser_press,
        tool_browser_screenshot,
        tool_browser_scroll,
        tool_browser_snapshot,
        tool_browser_type,
        tool_browser_wait,
    )

    return {
        "read": tool_read,
        "write": tool_write,
        "list": tool_list,
        "search": tool_search,
        "grep": tool_grep,
        "glob": tool_glob,
        "apply_patch": tool_apply_patch,
        "shell": tool_shell,
        "webfetch": tool_webfetch,
        "websearch": tool_websearch,
        "browser_open": tool_browser_open,
        "browser_snapshot": tool_browser_snapshot,
        "browser_click": tool_browser_click,
        "browser_type": tool_browser_type,
        "browser_wait": tool_browser_wait,
        "browser_scroll": tool_browser_scroll,
        "browser_press": tool_browser_press,
        "browser_screenshot": tool_browser_screenshot,
        "question": tool_question,
        "skill_load": tool_skill_load,
    }


def office_toolbelt() -> dict[str, Any]:
    """WorkBuddy 办公工具表：读写 + 成品落盘（无 shell）。"""
    return {
        "read": tool_read,
        "write": tool_write,
        "list": tool_list,
        "write_deliverable": tool_write_deliverable,
        "list_deliverables": tool_list_deliverables,
    }


_BUILTIN_IMPL: dict[str, Any] | None = None


def builtin_tool_impls() -> dict[str, Any]:
    global _BUILTIN_IMPL
    if _BUILTIN_IMPL is None:
        _BUILTIN_IMPL = {
            **coding_toolbelt(),
            **office_toolbelt(),
        }
    return _BUILTIN_IMPL


def reset_builtin_tool_impls_for_tests() -> None:
    global _BUILTIN_IMPL
    _BUILTIN_IMPL = None


def resolve_toolbelt(
    toolbelt: str | None,
    *,
    materials: dict[str, Any] | None = None,
    include_runtime: bool = False,
    bundle_root: str | Path | None = None,
) -> dict[str, Any]:
    """按原料注册表解析 toolbelt；可合并 materials.mcp 声明的工具。"""
    from fangyu.core.materials import load_materials, tool_ids_for_belt

    tb = (toolbelt or "coding").strip().lower()
    if materials is None:
        materials = load_materials(bundle_root) if bundle_root else None
        if materials is None:
            try:
                from fangyu.core.materials import default_materials
                materials = default_materials()
            except Exception:
                materials = None
    out: dict[str, Any] = {}
    if materials:
        ids = tool_ids_for_belt(tb, materials)
        impls = builtin_tool_impls()
        for tid in ids:
            if tid == "task" and not include_runtime:
                continue
            if tid in impls:
                out[tid] = impls[tid]
        # MCP 声明 → mcp_<tool> 可调用包装
        for mcp_tool_name, fn in _mcp_tools_from_materials(materials).items():
            out[mcp_tool_name] = fn
        if out:
            return out
    if tb == "office":
        return office_toolbelt()
    return coding_toolbelt()


def _mcp_tools_from_materials(materials: dict[str, Any]) -> dict[str, Any]:
    """materials.mcp: [{id, tools: [name,...]|\"*\"}] → 可调用工具。"""
    servers = materials.get("mcp") or []
    if not isinstance(servers, list):
        return {}
    out: dict[str, Any] = {}
    for srv in servers:
        if not isinstance(srv, dict):
            continue
        sid = str(srv.get("id") or "").strip() or "__internal__"
        names = srv.get("tools") or []
        if names == "*" or (isinstance(names, list) and len(names) == 1 and names[0] == "*"):
            names = _expand_mcp_tool_names(sid)
        if not isinstance(names, list):
            continue
        for raw in names:
            tname = str(raw).strip()
            if not tname or tname == "*":
                continue
            key = f"mcp_{tname}" if sid == "__internal__" else f"mcp_{sid}_{tname}"
            out[key] = _make_mcp_callable(sid, tname)
    return out


def _expand_mcp_tool_names(server: str) -> list[str]:
    """tools:\"*\" → 展开为可用工具名列表。"""
    if server == "__internal__":
        try:
            from fangyu.engine.tool_registry import list_tools, register_builtins
            register_builtins()
            return [
                str(t.get("name"))
                for t in list_tools()
                if t.get("enabled", True) and t.get("name")
            ]
        except Exception:
            return ["current_time"]
    try:
        from fangyu.engine.mcp import get_external_server
        conn = get_external_server(server)
        if not conn or not conn._tools_cache:
            return []
        return [str(t.get("name")) for t in conn._tools_cache if t.get("name")]
    except Exception:
        return []


def _make_mcp_callable(server: str, tool_name: str):
    async def _call(**kwargs: Any) -> Any:
        if server == "__internal__":
            from fangyu.engine.mcp import call_internal_tool, _init_internal_tools
            await _init_internal_tools()
            return await call_internal_tool(tool_name, kwargs or {})
        from fangyu.engine.mcp import get_external_server
        conn = get_external_server(server)
        if not conn:
            raise ValueError(f"MCP server 未连接: {server}")
        return await conn.call_tool(tool_name, kwargs or {})

    _call.__name__ = f"mcp_{tool_name}"
    _call.__doc__ = f"MCP {server}/{tool_name}"
    return _call


def _normalize_deliverable_rel(path: str, kind: str) -> str:
    rel = (path or "").strip().lstrip("/")
    if not rel:
        raise ValueError("deliverable path 为空")
    if ".." in Path(rel).parts:
        raise ValueError("非法路径")
    if not rel.startswith("deliverables/"):
        rel = f"deliverables/{rel}"
    p = Path(rel)
    kind_l = (kind or "md").lower().lstrip(".")
    if not p.suffix and kind_l:
        rel = str(p) + f".{kind_l}"
    return rel


def _minimal_docx_bytes(content: str) -> bytes:
    """零依赖：打包最小 OOXML docx（段落 + 简易标题/列表）。"""
    import zipfile
    from io import BytesIO
    from xml.sax.saxutils import escape

    def para(text: str, *, style: str | None = None) -> str:
        ppr = f'<w:pPr><w:pStyle w:val="{style}"/></w:pPr>' if style else ""
        return (
            f"<w:p>{ppr}<w:r><w:t xml:space=\"preserve\">"
            f"{escape(text)}</w:t></w:r></w:p>"
        )

    body_parts: list[str] = []
    for line in (content or "").replace("\r\n", "\n").split("\n"):
        raw = line.rstrip()
        if raw.startswith("# "):
            body_parts.append(para(raw[2:].strip() or "Untitled", style="Heading1"))
        elif raw.startswith("## "):
            body_parts.append(para(raw[3:].strip() or "Untitled", style="Heading2"))
        elif raw.startswith("### "):
            body_parts.append(para(raw[4:].strip() or "Untitled", style="Heading2"))
        elif raw.startswith("- ") or raw.startswith("* "):
            body_parts.append(para("• " + raw[2:].strip()))
        elif raw.strip():
            body_parts.append(para(raw))
    if not body_parts:
        body_parts.append(para(content or ""))

    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f"<w:body>{''.join(body_parts)}<w:sectPr/></w:body></w:document>"
    )
    content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>
"""
    rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>
"""
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", rels)
        zf.writestr("word/document.xml", document_xml)
    return buf.getvalue()


def _content_to_docx_bytes(content: str) -> bytes:
    """优先 python-docx；否则内置简易 OOXML。"""
    try:
        from docx import Document
        from io import BytesIO

        doc = Document()
        text = content or ""
        blocks = text.replace("\r\n", "\n").split("\n")
        para_buf: list[str] = []

        def flush_para() -> None:
            nonlocal para_buf
            if para_buf:
                doc.add_paragraph("\n".join(para_buf).strip())
                para_buf = []

        for line in blocks:
            raw = line.rstrip()
            if raw.startswith("# "):
                flush_para()
                doc.add_heading(raw[2:].strip() or "Untitled", level=1)
            elif raw.startswith("## "):
                flush_para()
                doc.add_heading(raw[3:].strip() or "Untitled", level=2)
            elif raw.startswith("### "):
                flush_para()
                doc.add_heading(raw[4:].strip() or "Untitled", level=3)
            elif raw.startswith("- ") or raw.startswith("* "):
                flush_para()
                doc.add_paragraph(raw[2:].strip(), style="List Bullet")
            elif not raw.strip():
                flush_para()
            else:
                para_buf.append(raw)
        flush_para()
        if not any(p.text.strip() for p in doc.paragraphs):
            doc.add_paragraph(text or "")
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError:
        return _minimal_docx_bytes(content)


def _col_name(idx: int) -> str:
    """0 -> A, 25 -> Z, 26 -> AA"""
    n = idx + 1
    letters = ""
    while n:
        n, r = divmod(n - 1, 26)
        letters = chr(65 + r) + letters
    return letters


def _parse_sheet_rows(content: str) -> list[list[str]]:
    """content → 二维表：JSON 数组 / Markdown 表 / CSV·TSV 行。"""
    text = (content or "").strip()
    if not text:
        return [[""]]
    # JSON
    if text.startswith("["):
        try:
            data = json.loads(text)
            if isinstance(data, list):
                rows: list[list[str]] = []
                for row in data:
                    if isinstance(row, list):
                        rows.append(["" if c is None else str(c) for c in row])
                    else:
                        rows.append([str(row)])
                return rows or [[""]]
        except json.JSONDecodeError:
            pass
    lines = text.replace("\r\n", "\n").split("\n")
    # Markdown table
    md_rows = [ln for ln in lines if ln.strip().startswith("|")]
    if len(md_rows) >= 2:
        rows = []
        for ln in md_rows:
            cells = [c.strip() for c in ln.strip().strip("|").split("|")]
            if cells and all(set(c) <= set("-: ") for c in cells):
                continue  # separator
            rows.append(cells)
        if rows:
            return rows
    # CSV / TSV
    rows = []
    for ln in lines:
        if not ln.strip():
            continue
        if "\t" in ln:
            rows.append([c.strip() for c in ln.split("\t")])
        else:
            rows.append([c.strip() for c in ln.split(",")])
    return rows or [[text]]


def _minimal_xlsx_bytes(content: str) -> bytes:
    """零依赖：最小 xlsx（一行一列表格）。"""
    import zipfile
    from io import BytesIO
    from xml.sax.saxutils import escape

    rows = _parse_sheet_rows(content)
    sheet_rows_xml: list[str] = []
    for r_i, row in enumerate(rows, start=1):
        cells = []
        for c_i, val in enumerate(row):
            ref = f"{_col_name(c_i)}{r_i}"
            # 数字尽量写 n，否则 inlineStr
            raw = val.strip()
            if raw and (raw.replace(".", "", 1).replace("-", "", 1).isdigit()):
                cells.append(f'<c r="{ref}"><v>{escape(raw)}</v></c>')
            else:
                cells.append(
                    f'<c r="{ref}" t="inlineStr"><is><t>{escape(val)}</t></is></c>'
                )
        sheet_rows_xml.append(f'<row r="{r_i}">{"".join(cells)}</row>')

    sheet_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
        f'<sheetData>{"".join(sheet_rows_xml)}</sheetData></worksheet>'
    )
    workbook_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main"
 xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <sheets><sheet name="Sheet1" sheetId="1" r:id="rId1"/></sheets>
</workbook>
"""
    content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>
  <Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>
</Types>
"""
    rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>
</Relationships>
"""
    wb_rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/>
</Relationships>
"""
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", rels)
        zf.writestr("xl/workbook.xml", workbook_xml)
        zf.writestr("xl/_rels/workbook.xml.rels", wb_rels)
        zf.writestr("xl/worksheets/sheet1.xml", sheet_xml)
    return buf.getvalue()


def tool_write_deliverable(path: str = "", content: str = "", kind: str = "md") -> str:
    """把成品写入 workspace/deliverables/（md / docx / xlsx）。"""
    kind_l = (kind or "md").lower().lstrip(".")
    probe = (path or "").strip()
    lower = probe.lower()
    if lower.endswith(".docx"):
        kind_l = "docx"
    elif lower.endswith(".xlsx"):
        kind_l = "xlsx"
    elif lower.endswith(".md") or lower.endswith(".txt") or lower.endswith(".csv"):
        kind_l = Path(probe).suffix.lstrip(".").lower() or kind_l

    rel = _normalize_deliverable_rel(path, kind_l)
    if kind_l == "docx":
        data = _content_to_docx_bytes(content)
        _ws().write_bytes(rel, data)
        return f"deliverable {rel} (docx, {len(data)} bytes)"
    if kind_l == "xlsx":
        data = _minimal_xlsx_bytes(content)
        _ws().write_bytes(rel, data)
        return f"deliverable {rel} (xlsx, {len(data)} bytes)"
    _ws().write(rel, content)
    return f"deliverable {rel} ({len(content)} chars)"


def tool_list_deliverables() -> list[str]:
    ws = _ws()
    root = ws.resolve("deliverables")
    if not root.exists():
        return []
    out: list[str] = []
    for f in sorted(root.rglob("*")):
        if f.is_file():
            out.append(str(f.relative_to(ws.root)))
    return out
